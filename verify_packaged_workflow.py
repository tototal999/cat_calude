"""Run the packaged document meeting workflow against a loopback fake LLM."""
from __future__ import annotations

import json
import os
import subprocess
import sys
import tempfile
import threading
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path


class _Handler(BaseHTTPRequestHandler):
    def do_POST(self):
        length = int(self.headers.get('Content-Length', '0'))
        payload = json.loads(self.rfile.read(length))
        prompt = payload['messages'][-1]['content']
        content = ('PACKAGED_MEETING_NOTES' if '產生可供會議使用的重點' in prompt
                   else 'PACKAGED_SUMMARY')
        body = json.dumps({
            'model': 'workflow-test-model',
            'choices': [{'message': {'content': content}}],
            'usage': {'total_tokens': 1},
        }).encode('utf-8')
        self.send_response(200)
        self.send_header('Content-Type', 'application/json')
        self.send_header('Content-Length', str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def log_message(self, _format, *_args):
        return


def main() -> int:
    executable = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('dist/ClaudeCat/ClaudeCat.exe')
    if not executable.is_file():
        print(f'QA_RESULT|STATUS:FAIL|EXPECTED:packaged executable|ACTUAL:missing {executable}')
        return 1

    from docx import Document

    server = ThreadingHTTPServer(('127.0.0.1', 0), _Handler)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    try:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / 'meeting.docx'
            document = Document()
            document.add_heading('Payment terms', level=1)
            document.add_paragraph('Payment is due in 30 days.')
            document.save(source)
            environment = os.environ.copy()
            environment['LOCALAPPDATA'] = str(root / 'appdata')
            endpoint = f'http://127.0.0.1:{server.server_port}/v1'
            completed = subprocess.run(
                [str(executable), '--workflow-check', str(source), endpoint],
                timeout=120,
                env=environment,
            )
            if completed.returncode != 0:
                print('QA_RESULT|STATUS:FAIL|EXPECTED:completed meeting-pack workflow'
                      f'|ACTUAL:exit {completed.returncode}')
                return 1
    finally:
        server.shutdown()
        server.server_close()
    print('QA_RESULT|STATUS:PASS|EXPECTED:DOCX -> evidence -> LLM -> Markdown artifact'
          '|ACTUAL:packaged workflow completed')
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
